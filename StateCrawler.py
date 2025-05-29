#!/usr/bin/env python3
"""
STATE AUTISM EMPLOYMENT PROGRAMS CRAWLER
========================================

This crawler specifically targets state-level programs that assist 
people with Autism Spectrum Disorder (ASD) in finding employment.

Prerequisites:
pip install requests beautifulsoup4 pandas python-docx
"""

import os
import csv
import time
import queue
import logging
import re
from urllib.parse import urljoin, urlparse
from urllib.robotparser import RobotFileParser

import requests
from bs4 import BeautifulSoup
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# Optional imports with fallbacks
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not available. Word report generation will be skipped.")
    DOCX_AVAILABLE = False

# Setup logging and output directory
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Create output directories
output_dir = "StateAutismEmploymentPrograms"
txt_dir = os.path.join(output_dir, "programs")
os.makedirs(txt_dir, exist_ok=True)
csv_path = os.path.join(output_dir, "state_programs_index.csv")
report_path = os.path.join(output_dir, "StateAutismEmploymentReport.docx")

# US States for pattern matching
US_STATES = [
    'alabama', 'alaska', 'arizona', 'arkansas', 'california', 'colorado',
    'connecticut', 'delaware', 'florida', 'georgia', 'hawaii', 'idaho',
    'illinois', 'indiana', 'iowa', 'kansas', 'kentucky', 'louisiana',
    'maine', 'maryland', 'massachusetts', 'michigan', 'minnesota',
    'mississippi', 'missouri', 'montana', 'nebraska', 'nevada',
    'new hampshire', 'new jersey', 'new mexico', 'new york',
    'north carolina', 'north dakota', 'ohio', 'oklahoma', 'oregon',
    'pennsylvania', 'rhode island', 'south carolina', 'south dakota',
    'tennessee', 'texas', 'utah', 'vermont', 'virginia', 'washington',
    'west virginia', 'wisconsin', 'wyoming'
]

# State abbreviations
STATE_ABBREVS = [
    'AL', 'AK', 'AZ', 'AR', 'CA', 'CO', 'CT', 'DE', 'FL', 'GA',
    'HI', 'ID', 'IL', 'IN', 'IA', 'KS', 'KY', 'LA', 'ME', 'MD',
    'MA', 'MI', 'MN', 'MS', 'MO', 'MT', 'NE', 'NV', 'NH', 'NJ',
    'NM', 'NY', 'NC', 'ND', 'OH', 'OK', 'OR', 'PA', 'RI', 'SC',
    'SD', 'TN', 'TX', 'UT', 'VT', 'VA', 'WA', 'WV', 'WI', 'WY'
]

class StateAutismEmploymentCrawler:
    def __init__(self, delay=2.0, max_retries=3):
        self.delay = delay
        self.session = requests.Session()
        
        # Setup retry strategy
        retry_strategy = Retry(
            total=max_retries,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)
        
        # Set headers to appear as a regular browser
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def can_fetch(self, url):
        """Check robots.txt compliance"""
        try:
            parsed_url = urlparse(url)
            robots_url = f"{parsed_url.scheme}://{parsed_url.netloc}/robots.txt"
            
            rp = RobotFileParser()
            rp.set_url(robots_url)
            rp.read()
            return rp.can_fetch('*', url)
        except Exception:
            return True

    def get_page_content(self, url):
        """Fetch page content with error handling"""
        try:
            if not self.can_fetch(url):
                logger.info(f"Robots.txt disallows crawling: {url}")
                return None
                
            response = self.session.get(url, timeout=15)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '').lower()
            if 'text/html' not in content_type:
                logger.info(f"Skipping non-HTML content: {url}")
                return None
                
            return response.text
            
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to fetch {url}: {e}")
            return None

    def extract_links(self, html, base_url):
        """Extract all links from HTML content"""
        try:
            soup = BeautifulSoup(html, 'html.parser')
            links = set()
            
            for tag in soup.find_all('a', href=True):
                href = tag.get('href')
                if href:
                    abs_url = urljoin(base_url, href)
                    if abs_url.startswith(('http://', 'https://')):
                        links.add(abs_url)
            
            return links
        except Exception as e:
            logger.error(f"Error extracting links from {base_url}: {e}")
            return set()

    def extract_state_from_content(self, text, url):
        """Extract state information from content and URL"""
        text_lower = text.lower()
        url_lower = url.lower()
        
        # Check URL for state indicators
        for state in US_STATES:
            if state in url_lower:
                return state.title()
        
        for abbrev in STATE_ABBREVS:
            if f"/{abbrev.lower()}/" in url_lower or f".{abbrev.lower()}." in url_lower:
                return abbrev
        
        # Check content for state mentions
        state_pattern = r'\b(' + '|'.join(US_STATES + [s.lower() for s in STATE_ABBREVS]) + r')\b'
        matches = re.findall(state_pattern, text_lower)
        
        if matches:
            # Return most frequently mentioned state
            from collections import Counter
            most_common = Counter(matches).most_common(1)
            return most_common[0][0].title() if most_common else "Unknown"
        
        return "Unknown"

    def extract_program_details(self, soup, text):
        """Extract specific program details"""
        details = {
            'program_name': '',
            'services': [],
            'eligibility': '',
            'contact_info': '',
            'funding_source': ''
        }
        
        # Extract program name from title or main heading
        if soup.title:
            details['program_name'] = soup.title.string.strip()
        else:
            h1 = soup.find('h1')
            if h1:
                details['program_name'] = h1.get_text().strip()
        
        # Look for service descriptions
        service_keywords = [
            'job training', 'employment support', 'vocational rehabilitation',
            'career counseling', 'job placement', 'supported employment',
            'workplace accommodation', 'job coaching', 'skills training',
            'internship', 'apprenticeship', 'transition services'
        ]
        
        text_lower = text.lower()
        for keyword in service_keywords:
            if keyword in text_lower:
                details['services'].append(keyword)
        
        # Look for contact information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
        
        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        
        contact_parts = []
        if emails:
            contact_parts.append(f"Email: {emails[0]}")
        if phones:
            contact_parts.append(f"Phone: {phones[0]}")
        
        details['contact_info'] = "; ".join(contact_parts)
        
        # Look for eligibility criteria
        eligibility_indicators = ['eligible', 'qualification', 'requirement', 'criteria']
        sentences = text.split('.')
        eligibility_sentences = []
        
        for sentence in sentences:
            if any(indicator in sentence.lower() for indicator in eligibility_indicators):
                if any(keyword in sentence.lower() for keyword in ['autism', 'disability', 'developmental']):
                    eligibility_sentences.append(sentence.strip())
        
        details['eligibility'] = ". ".join(eligibility_sentences[:2])  # Limit to 2 sentences
        
        # Look for funding information
        funding_keywords = ['funded by', 'grant', 'federal', 'state funding', 'department of']
        for sentence in sentences:
            if any(keyword in sentence.lower() for keyword in funding_keywords):
                details['funding_source'] = sentence.strip()
                break
        
        return details

    def is_state_autism_employment_content(self, text, url):
        """Check if content is relevant to state autism employment programs"""
        text_lower = text.lower()
        url_lower = url.lower()
        
        # Must have autism-related terms
        autism_terms = ['autism', 'autistic', 'asd', 'asperger', 'developmental disability']
        has_autism = any(term in text_lower for term in autism_terms)
        
        # Must have employment-related terms
        employment_terms = [
            'employment', 'job', 'work', 'career', 'vocational', 'workplace',
            'hiring', 'training', 'placement', 'supported employment'
        ]
        has_employment = any(term in text_lower for term in employment_terms)
        
        # Must have state/government indicators
        state_indicators = [
            'state', 'government', 'department', 'division', 'agency',
            'public', 'gov', '.gov', 'rehabilitation', 'services'
        ]
        has_state_indicator = any(indicator in text_lower or indicator in url_lower 
                                for indicator in state_indicators)
        
        # Must have program indicators
        program_indicators = [
            'program', 'service', 'initiative', 'support', 'assistance',
            'help', 'resource', 'opportunity'
        ]
        has_program = any(indicator in text_lower for indicator in program_indicators)
        
        # Check for state names in URL or content
        has_state_name = (any(state in url_lower for state in US_STATES) or 
                         any(state in text_lower for state in US_STATES) or
                         any(abbrev.lower() in url_lower for abbrev in STATE_ABBREVS))
        
        relevance_score = sum([has_autism, has_employment, has_state_indicator, has_program, has_state_name])
        
        return relevance_score >= 3  # Require at least 3 out of 5 criteria

    def crawl(self, start_urls, max_pages=100, max_depth=3):
        """Main crawling function focused on state programs"""
        visited = set()
        results = []
        q = queue.Queue()
        
        # Initialize queue with start URLs
        for url in start_urls:
            q.put((url, 0))
        
        logger.info(f"Starting state autism employment program crawl with {len(start_urls)} seed URLs")
        
        while not q.empty() and len(visited) < max_pages:
            url, depth = q.get()
            
            if url in visited or depth > max_depth:
                continue
                
            visited.add(url)
            
            # Rate limiting
            time.sleep(self.delay)
            
            html = self.get_page_content(url)
            if not html:
                continue
            
            try:
                soup = BeautifulSoup(html, 'html.parser')
                
                # Remove unwanted elements
                for script in soup(["script", "style", "nav", "footer", "header"]):
                    script.decompose()
                
                text = soup.get_text(separator='\n', strip=True)
                
                # Check relevance to state autism employment
                if not self.is_state_autism_employment_content(text, url):
                    continue
                
                # Extract title
                title = ""
                if soup.title:
                    title = soup.title.string.strip()
                
                # Extract state information
                state = self.extract_state_from_content(text, url)
                
                # Extract program details
                program_details = self.extract_program_details(soup, text)
                
                # Extract domain/organization
                domain = urlparse(url).netloc.replace("www.", "")
                
                result = {
                    'url': url,
                    'title': title,
                    'state': state,
                    'domain': domain,
                    'text': text,
                    'depth': depth,
                    'program_name': program_details['program_name'],
                    'services': ', '.join(program_details['services']),
                    'eligibility': program_details['eligibility'],
                    'contact_info': program_details['contact_info'],
                    'funding_source': program_details['funding_source']
                }
                
                results.append(result)
                
                logger.info(f"Found state program ({len(results)}/{max_pages}): {state} - {title[:60]}...")
                
                # Add new links to queue if we haven't reached max depth
                if depth < max_depth:
                    try:
                        links = self.extract_links(html, url)
                        for link in links:
                            if link not in visited and len(visited) + q.qsize() < max_pages * 2:
                                q.put((link, depth + 1))
                    except Exception as e:
                        logger.error(f"Error processing links from {url}: {e}")
                        
            except Exception as e:
                logger.error(f"Error processing {url}: {e}")
                continue
        
        logger.info(f"Crawling completed. Found {len(results)} state autism employment programs.")
        return results

def save_results(results):
    """Save results to files with state program focus"""
    
    # Save individual text files and create CSV index
    with open(csv_path, mode='w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow([
            "URL", "Title", "State", "Organization", "Program_Name", 
            "Services", "Eligibility", "Contact_Info", "Funding_Source", "Filename"
        ])
        
        for i, item in enumerate(results, 1):
            filename = f"state_program_{i}.txt"
            filepath = os.path.join(txt_dir, filename)
            
            # Save detailed text content
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"STATE: {item['state']}\n")
                f.write(f"PROGRAM: {item['program_name']}\n")
                f.write(f"URL: {item['url']}\n")
                f.write(f"ORGANIZATION: {item['domain']}\n")
                f.write(f"SERVICES: {item['services']}\n")
                f.write(f"ELIGIBILITY: {item['eligibility']}\n")
                f.write(f"CONTACT: {item['contact_info']}\n")
                f.write(f"FUNDING: {item['funding_source']}\n")
                f.write("-" * 70 + "\n")
                f.write(item['text'])
            
            writer.writerow([
                item['url'], item['title'], item['state'], item['domain'],
                item['program_name'], item['services'], item['eligibility'],
                item['contact_info'], item['funding_source'], filename
            ])
    
    logger.info(f"Saved {len(results)} state program files and CSV index")
    
    # Generate summary statistics by state
    generate_state_summary(results)
    
    # Generate Word report if available
    if DOCX_AVAILABLE and results:
        generate_state_report(results)

def generate_state_summary(results):
    """Generate summary statistics by state"""
    from collections import defaultdict, Counter
    
    # Group by state
    state_programs = defaultdict(list)
    for result in results:
        state_programs[result['state']].append(result)
    
    # Create summary file
    summary_path = os.path.join(output_dir, "state_summary.txt")
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write("STATE AUTISM EMPLOYMENT PROGRAMS SUMMARY\n")
        f.write("="*50 + "\n\n")
        
        f.write(f"Total Programs Found: {len(results)}\n")
        f.write(f"States Represented: {len(state_programs)}\n\n")
        
        # Sort states by number of programs
        sorted_states = sorted(state_programs.items(), key=lambda x: len(x[1]), reverse=True)
        
        for state, programs in sorted_states:
            f.write(f"\n{state.upper()}: {len(programs)} programs\n")
            f.write("-" * 30 + "\n")
            
            for program in programs:
                f.write(f"• {program['program_name']}\n")
                if program['services']:
                    f.write(f"  Services: {program['services']}\n")
                if program['contact_info']:
                    f.write(f"  Contact: {program['contact_info']}\n")
                f.write(f"  URL: {program['url']}\n\n")
    
    logger.info(f"State summary saved to {summary_path}")

def generate_state_report(results):
    """Generate Word document report focused on state programs"""
    try:
        doc = Document()
        doc.add_heading("State Autism Employment Programs Report", level=1)
        
        # Introduction
        intro = doc.add_paragraph()
        intro.add_run(f"This report catalogs {len(results)} state-level programs designed to assist individuals with Autism Spectrum Disorder (ASD) in finding and maintaining employment. ")
        intro.add_run("The programs are organized by state to facilitate easy reference for individuals, families, and service providers.")
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=2)
        
        from collections import defaultdict
        state_counts = defaultdict(int)
        all_services = []
        
        for result in results:
            state_counts[result['state']] += 1
            if result['services']:
                all_services.extend(result['services'].split(', '))
        
        doc.add_paragraph(f"• Total programs identified: {len(results)}")
        doc.add_paragraph(f"• States with programs: {len(state_counts)}")
        
        # Top states by number of programs
        top_states = sorted(state_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        top_states_text = ", ".join([f"{state} ({count})" for state, count in top_states])
        doc.add_paragraph(f"• States with most programs: {top_states_text}")
        
        # Most common services
        from collections import Counter
        service_counts = Counter(all_services)
        top_services = service_counts.most_common(5)
        if top_services:
            top_services_text = ", ".join([service for service, count in top_services])
            doc.add_paragraph(f"• Most common services: {top_services_text}")
        
        # Programs by State
        doc.add_heading("Programs by State", level=2)
        
        # Group and sort by state
        state_programs = defaultdict(list)
        for result in results:
            state_programs[result['state']].append(result)
        
        for state in sorted(state_programs.keys()):
            programs = state_programs[state]
            doc.add_heading(f"{state} ({len(programs)} programs)", level=3)
            
            for i, program in enumerate(programs, 1):
                doc.add_heading(f"{i}. {program['program_name']}", level=4)
                
                if program['services']:
                    doc.add_paragraph(f"Services: {program['services']}")
                
                if program['eligibility']:
                    doc.add_paragraph(f"Eligibility: {program['eligibility']}")
                
                if program['contact_info']:
                    doc.add_paragraph(f"Contact: {program['contact_info']}")
                
                if program['funding_source']:
                    doc.add_paragraph(f"Funding: {program['funding_source']}")
                
                doc.add_paragraph(f"Website: {program['url']}")
                doc.add_paragraph("")  # Add spacing
        
        doc.save(report_path)
        logger.info(f"State programs report saved to {report_path}")
        
    except Exception as e:
        logger.error(f"Failed to generate Word report: {e}")

def main():
    """Main execution function"""
    
    # Seed URLs focused on state government and disability services
    seed_urls = [
        # Federal resources that link to state programs
        "https://www.dol.gov/agencies/odep/",
        "https://www.acl.gov/programs/aging-and-disability-networks/state-disability",
        "https://www.ed.gov/about/offices/list/osers/rsa/",
        
        # Autism organizations with state program directories
        "https://www.autismspeaks.org/state-autism-profiles",
        "https://www.autism-society.org/living-with-autism/",
        
        # Vocational rehabilitation state directories
        "https://www.csavr.org/content/members",
        "https://www.nchrtm.org/",
        
        # Developmental disabilities councils
        "https://nacdd.org/",
        "https://www.aucd.org/",
        
        # Employment focused
        "https://www.thinkwork.org/",
        "https://www.apse.org/",
        
        # Specific high-population state examples
        "https://www.dor.ca.gov/",  # California
        "https://www.twc.texas.gov/",  # Texas
        "https://www.acces.nysed.gov/vr/",  # New York
        "https://www.fldoe.org/academics/career-adult-edu/vocational-rehabilitation/",  # Florida
    ]
    
    logger.info("Starting State Autism Employment Programs Crawler")
    
    # Initialize crawler
    crawler = StateAutismEmploymentCrawler(delay=3.0)  # 3 second delay for government sites
    
    # Perform crawling
    results = crawler.crawl(seed_urls, max_pages=75, max_depth=3)
    
    if not results:
        logger.error("No state autism employment programs found. Try expanding search criteria.")
        return
    
    # Save results
    save_results(results)
    
    logger.info("State autism employment program crawling completed successfully!")
    logger.info(f"Results saved in '{output_dir}' directory")
    logger.info(f"Check 'state_summary.txt' for a quick overview by state")

if __name__ == "__main__":
    main()