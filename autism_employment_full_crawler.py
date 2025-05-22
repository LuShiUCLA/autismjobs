
# ============================================================================
# AUTISM EMPLOYMENT PROGRAM CRAWLER & NLP ANALYZER (Python Companion Version)
# ============================================================================

# --- 1. Install Required Packages ---
# pip install requests beautifulsoup4 selenium pandas spacy python-docx transformers

import os
import csv
import time
import queue
import logging
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup

import spacy
from transformers import pipeline
from docx import Document

from selenium import webdriver
from selenium.webdriver.chrome.options import Options

# --- Setup Logging and Output Directory ---
logging.basicConfig(level=logging.INFO)
output_dir = "C:/data/AutismEmployment"
txt_dir = os.path.join(output_dir, "pages")
os.makedirs(txt_dir, exist_ok=True)
csv_path = os.path.join(output_dir, "index.csv")
report_path = os.path.join(output_dir, "AutismEmploymentReport.docx")

# --- Load NLP Models ---
nlp = spacy.load("en_core_web_sm")
summarizer = pipeline("summarization")
sentiment_analyzer = pipeline("sentiment-analysis")

# --- Fetch Page Content ---
def get_page_content(url, use_selenium=False):
    try:
        if use_selenium:
            options = Options()
            options.add_argument('--headless')
            driver = webdriver.Chrome(options=options)
            driver.set_page_load_timeout(15)
            driver.get(url)
            html = driver.page_source
            driver.quit()
        else:
            resp = requests.get(url, timeout=10)
            resp.raise_for_status()
            html = resp.text
        return html
    except Exception as e:
        logging.error(f"Failed to fetch {url}: {e}")
        return None

# --- Extract Links ---
def extract_links(html, base_url):
    soup = BeautifulSoup(html, 'html.parser')
    links = set()
    for tag in soup.find_all('a', href=True):
        href = tag.get('href')
        abs_url = urljoin(base_url, href)
        if abs_url.startswith("http"):
            links.add(abs_url)
    return links

# --- Main Crawler ---
def crawl(start_urls, max_pages=200, max_depth=3):
    visited = set()
    results = []
    q = queue.Queue()
    for url in start_urls:
        q.put((url, 0))

    while not q.empty() and len(visited) < max_pages:
        url, depth = q.get()
        if url in visited or depth > max_depth:
            continue
        visited.add(url)

        html = get_page_content(url)
        if not html:
            continue

        soup = BeautifulSoup(html, 'html.parser')
        title = soup.title.string.strip() if soup.title else ''
        text = soup.get_text(separator='\n', strip=True)

        keywords = ["autism", "employment", "program", "effectiveness", "outcome", "evaluation"]
        if not any(kw in text.lower() for kw in keywords):
            continue

        domain = urlparse(url).netloc.replace("www.", "")
        results.append((url, title, domain, text))
        logging.info(f"Crawled: {url} (depth {depth})")

        if depth < max_depth:
            for link in extract_links(html, url):
                if link not in visited:
                    q.put((link, depth + 1))
    return results

# --- Seed URLs ---
seed_urls = [
    "https://www.autismspeaks.org/family-services/autism-work",
    "https://www.autism.org/autism-services",
    "https://www.dol.gov/agencies/odep",
    "https://www.spectrumworks.org",
    "https://www.nimh.nih.gov/health/topics/autism-spectrum-disorders-asd",
]

results = crawl(seed_urls)

# --- Save .txt and CSV Index ---
with open(csv_path, mode='w', newline='', encoding='utf-8') as csvfile:
    writer = csv.writer(csvfile)
    writer.writerow(["URL", "Title", "Organization", "Filename"])
    for i, (url, title, org, text) in enumerate(results, 1):
        filename = f"page_{i}.txt"
        filepath = os.path.join(txt_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(text)
        writer.writerow([url, title, org, filename])

# --- NLP Analysis ---
analysis_results = []
for url, title, org, text in results:
    try:
        summary = summarizer(text[:1000], max_length=130, min_length=30, do_sample=False)[0]['summary_text']
    except Exception:
        summary = ""

    doc = nlp(text)
    entities = [(ent.text, ent.label_) for ent in doc.ents]

    try:
        sentiment = sentiment_analyzer(text[:512])
    except Exception:
        sentiment = []

    analysis_results.append({
        "url": url, "title": title, "organization": org,
        "summary": summary, "entities": entities, "sentiment": sentiment
    })

# --- Word Report Generation ---
doc = Document()
doc.add_heading("Autism Employment Programs: Evidence and Outcomes", level=1)
doc.add_paragraph("This report summarizes web-sourced programs designed to improve employment outcomes for autistic individuals.")

for i, item in enumerate(analysis_results, 1):
    doc.add_heading(f"{i}. {item['organization']}", level=2)
    doc.add_paragraph(item['title'], style='List Bullet')
    doc.add_paragraph(f"Source: {item['url']}")
    doc.add_heading("Summary", level=3)
    doc.add_paragraph(item['summary'] or "No summary generated.")
    doc.add_heading("Named Entities", level=3)
    ents = ", ".join([f"{txt} ({lbl})" for txt, lbl in item['entities']]) or "None"
    doc.add_paragraph(ents)
    doc.add_heading("Sentiment", level=3)
    sentiments = "; ".join([f"{s['label']} ({s['score']:.2f})" for s in item['sentiment']]) or "Not analyzed"
    doc.add_paragraph(sentiments)

doc.add_heading("References", level=1)
for i, item in enumerate(analysis_results, 1):
    doc.add_paragraph(f"[{i}] {item['title']}. Available at: {item['url']}")

doc.save(report_path)
logging.info(f"✔ Word report saved to {report_path}")
